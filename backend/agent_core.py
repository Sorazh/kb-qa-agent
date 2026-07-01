"""
知识库问答 Agent (LangChain v1)

支持 txt/pdf/docx 文档加载、知识库检索与问答，对话记忆由 MemorySaver 管理。
"""

import warnings

from database import clear_all_records

warnings.filterwarnings("ignore", category=DeprecationWarning)
warnings.filterwarnings("ignore", message=".*unauthenticated requests to the HF Hub.*")
import os
import json
from typing import Any, List
from typing import NotRequired
from dotenv import load_dotenv

load_dotenv()

from langchain_openai import ChatOpenAI
from langchain.tools import tool
from langchain.agents import create_agent, AgentState
from langchain.agents.middleware import AgentMiddleware
from langchain.messages import AIMessage, HumanMessage
from langgraph.checkpoint.memory import MemorySaver
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

import asyncio
from database import get_all_files, add_file_record, delete_file_record,clear_all_records

# ===== 配置 =====
FAISS_INDEX_DIR = "./faiss_index"
EMBED_MODEL_NAME = "all-MiniLM-L6-v2"

_knowledge_base: FAISS | None = None


model = ChatOpenAI(
    model="deepseek-v4-flash",
    api_key=os.getenv("DEEPSEEK_API_KEY"),
    base_url=os.getenv("DEEPSEEK_BASE_URL"),
    temperature=0.7,
)


def run_query(user_input: str, thread_id: str = "default") -> str:
    """调用 Agent 回答问题，返回文本结果。"""
    config = {"configurable": {"thread_id": thread_id}}
    result = agent.invoke(
        {"messages": [HumanMessage(content=user_input)]},
        config=config,
    )
    for msg in reversed(result["messages"]):
        if isinstance(msg, AIMessage):
            return msg.content
    return "(无响应)"


def _get_embeddings():
    return HuggingFaceEmbeddings(model_name=EMBED_MODEL_NAME)


def _ensure_kb_loaded():
    """从磁盘恢复知识库（如果持久化索引存在）。"""
    global _knowledge_base
    if _knowledge_base is not None:
        return
    if os.path.isdir(FAISS_INDEX_DIR) and os.path.exists(os.path.join(FAISS_INDEX_DIR, "index.faiss")):
        _knowledge_base = FAISS.load_local(
            FAISS_INDEX_DIR,
            _get_embeddings(),
            allow_dangerous_deserialization=True,
        )
        print(f"[KB] 已从磁盘恢复知识库：{FAISS_INDEX_DIR}")
    else:
        print("[KB] 未检测到持久化知识库，等待首次 load_document...")


def _save_kb():
    global _knowledge_base
    if _knowledge_base is None:
        return
    _knowledge_base.save_local(FAISS_INDEX_DIR)
    print(f"[KB] 知识库已持久化到：{FAISS_INDEX_DIR}")

def _list_loaded_files() -> list[str]:
    """从数据库获取所有已加载的文件路径"""
    try:
        # asyncio.run() 用来在同步函数里跑异步代码
        # 这里调用了 database.py 里的 get_all_files()
        return asyncio.run(get_all_files())
    except Exception as e:
        print(f"[DB] 读取文件列表失败：{e}")
        return []  #如果读不到，返回空列表，不影响程序运行

def _record_loaded_file(filepath: str,chunk_count: int = 0):
    """记录一个文件到数据库（同步版本）"""
    try:
        asyncio.run(add_file_record(filepath, chunk_count))
    except Exception as e:
        print(f"[DB] 记录文件失败：{e}")

def _delete_file_record(filepath: str):
    """同步包装器：从数据库删除文件"""
    try:
        asyncio.run(delete_file_record(filepath))
    except Exception as e:
        print(f"删除数据库记录失败: {e}")

def _load_file_to_documents(filepath: str) -> List[Document]:
    """根据扩展名加载文件，返回 Document 列表。"""
    ext = os.path.splitext(filepath)[1].lower()
    documents = []

    if ext == ".txt":
        with open(filepath, "r", encoding="utf-8") as f:
            text = f.read()
        documents.append(Document(page_content=text, metadata={"source": filepath}))

    elif ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        documents.append(Document(page_content=text, metadata={"source": filepath}))

    elif ext == ".docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(filepath)
        text = "\n".join([para.text for para in doc.paragraphs])
        documents.append(Document(page_content=text, metadata={"source": filepath}))

    else:
        raise ValueError(f"不支持的文件格式: {ext}，仅支持 .txt .pdf .docx")

    return documents

def do_load_document(filepath: str) -> str:
    """加载文档到知识库：读取 → 切分 → 向量化 → 持久化。"""
    try:
        if not os.path.exists(filepath):
            return f"文件不存在: {filepath}"

        _ensure_kb_loaded()

        docs = _load_file_to_documents(filepath)
        #经验证：chunk_size=300, overlap=30 适合中文短句检索
        #未来优化方向：根据文档类型（技术文档/小说）动态调整分块策略
        splitter = RecursiveCharacterTextSplitter(chunk_size=300, chunk_overlap=30)
        chunks = splitter.split_documents(docs)

        global _knowledge_base
        embeddings = _get_embeddings()

        if _knowledge_base is None:
            _knowledge_base = FAISS.from_documents(chunks, embeddings)
        else:
            _knowledge_base.add_documents(chunks)

        _save_kb()
        _record_loaded_file(filepath)

        return (
            f"成功加载文档: {filepath}\n"
            f"本次新增片段数: {len(chunks)}\n"
            f"知识库已持久化到: {FAISS_INDEX_DIR}"
        )
    except Exception as e:
        return f"加载文档失败: {str(e)}"


def do_remove_document(filepath: str) -> str:
    """从知识库删除指定文档，剩余文件重建索引。"""
    try:
        abs_path = os.path.abspath(filepath)
        current_files = _list_loaded_files()
        if not current_files:
            return "当前知识库中没有记录任何文件，无法删除。"

        matched = None
        for f in current_files:
            if os.path.normcase(f) == os.path.normcase(abs_path):
                matched = f
                break

        if matched is None:
            return (f"未找到已导入的文件：{filepath}\n"
                    f"当前已导入的文件：\n" + "\n".join(current_files))

        new_files = [f for f in current_files if f != matched]

        rebuild_knowledge_base_from_files(new_files)
        # 从数据库里删除这条记录（不删的话文件列表里还会显示）
        _delete_file_record(matched)

        return (f"✅ 已成功删除文档：{matched}\n"
                f"知识库已重建并持久化。当前共 {len(new_files)} 个文件。")
    except Exception as e:
        return f"删除文档失败: {str(e)}"


@tool
def load_document(filepath: str) -> str:
    """加载一个文档文件（txt/pdf/docx）到知识库中，供后续查询使用。"""
    return do_load_document(filepath)


@tool
def query_knowledge(query: str) -> str:
    """从知识库中检索与查询最相关的信息并返回。"""
    _ensure_kb_loaded()

    global _knowledge_base
    if _knowledge_base is None:
        loaded = _list_loaded_files()
        hint = f"\n（已记录的导入文件：{loaded}）" if loaded else ""
        return "知识库为空，请先使用 load_document 工具加载文档。" + hint

    results = _knowledge_base.similarity_search(query, k=3)
    if not results:
        return "未找到相关信息。"

    combined = "\n---\n".join([doc.page_content for doc in results])
    return f"找到以下相关内容：\n{combined}"


def rebuild_knowledge_base_from_files(filepaths: List[str]):
    """根据文件列表重新构建整个知识库（清空旧索引）。"""
    global _knowledge_base
    _knowledge_base = None

    if not filepaths:
        import shutil
        if os.path.isdir(FAISS_INDEX_DIR):
            shutil.rmtree(FAISS_INDEX_DIR)
        try:
            asyncio.run(clear_all_records())
            print(f"[KB] 所有文件已删除，数据库记录已同步清空")
        except Exception as e:
            print(f"[KB] 清空数据库记录失败：{e}")
        return

    all_chunks = []
    embeddings = _get_embeddings()
    splitter = RecursiveCharacterTextSplitter(chunk_size=300, chunk_overlap=30)

    for fp in filepaths:
        if not os.path.exists(fp):
            continue
        docs = _load_file_to_documents(fp)
        chunks = splitter.split_documents(docs)
        all_chunks.extend(chunks)

    if all_chunks:
        _knowledge_base = FAISS.from_documents(all_chunks, embeddings)
        _save_kb()
    else:
        _knowledge_base = None


@tool
def remove_document(filepath: str) -> str:
    """从知识库中删除一个已导入的文档，知识库会自动重建。"""
    return do_remove_document(filepath)


class MyAgentState(AgentState):
    user_name: NotRequired[str | None]


class LoggingMiddleware(AgentMiddleware):
    def before_model(self, state: MyAgentState, runtime) -> dict[str, Any] | None:
        print(f"[中间件] 模型调用前，消息数：{len(state['messages'])}")
        return None

    def after_model(self, state: MyAgentState, runtime) -> dict[str, Any] | None:
        last = state["messages"][-1] if state["messages"] else None
        if isinstance(last, AIMessage):
            print(f"[中间件] 模型响应长度：{len(last.content)} 字符")
        return None


agent = create_agent(
    model=model,
    tools=[load_document, query_knowledge, remove_document],
    system_prompt=(
        "你是一个知识库问答助手。你的工作流程如下：\n\n"
        "1. 当用户提供文件路径（以.txt、.pdf 或.docx 结尾）时，"
        "立即调用 load_document 工具加载该文件。\n"
        "2. 当用户提出问题时，**必须首先调用 query_knowledge 工具**尝试检索知识库。\n"
        "   - 如果检索结果不为空，则基于检索结果回答。\n"
        "   - 如果检索结果为空，再询问用户是否需要加载文档。\n"
        "3. 知识库支持磁盘持久化，重启后会自动恢复。\n"
        "4. 当用户要求删除某个已导入的文档时，调用 remove_document 工具。\n"
        "5. 不要在未调用工具的情况下凭空回答知识库相关的问题。\n"
    ),
    state_schema=MyAgentState,
    middleware=[LoggingMiddleware()],
    checkpointer=MemorySaver(),
)


if __name__ == "__main__":
    print("=== 知识库问答 Agent (LangChain v1) ===")
    existing_files = _list_loaded_files()
    if existing_files:
        print("上次已导入的文件：")
        for f in existing_files:
            print(f"  - {f}")

    thread_id = input("请输入会话ID（留空自动生成）: ") or "default-thread"
    config = {"configurable": {"thread_id": thread_id}}

    while True:
        user_input = input("\n你: ").strip()
        if user_input.lower() in ("exit", "quit", "拜拜"):
            print("Agent: 再见！")
            break

        result = agent.invoke(
            {"messages": [{"role": "user", "content": user_input}]},
            config=config,
        )
        for msg in reversed(result["messages"]):
            if isinstance(msg, AIMessage):
                print(f"Agent: {msg.content}")
                break


